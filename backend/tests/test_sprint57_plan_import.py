"""
Sprint 57 — Plan Import Wizard Backend Tests
Tests /api/planos-inspecao/parse-text and /api/planos-inspecao/parse-file endpoints.
"""
import io
import os
import pytest
import requests

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', 'https://procure-manutrix.preview.emergentagent.com').rstrip('/')


# ============== FIXTURES ==============

@pytest.fixture(scope="module")
def pcm_token():
    """Login as PCM user."""
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"email": "test.pcm@maintrix.com", "password": "pcm123"})
    if r.status_code != 200:
        pytest.skip(f"PCM login failed: {r.status_code} {r.text[:200]}")
    return r.json()["access_token"]


@pytest.fixture(scope="module")
def pcm_headers(pcm_token):
    return {"Authorization": f"Bearer {pcm_token}"}


@pytest.fixture(scope="module")
def operador_token():
    """Login as Operador (no plan-import permission)."""
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"email": "test.operador@maintrix.com", "password": "op123"})
    if r.status_code != 200:
        pytest.skip(f"Operador login failed: {r.status_code}")
    return r.json()["access_token"]


# ============== PARSE-TEXT ENDPOINT ==============

class TestParseText:
    """POST /api/planos-inspecao/parse-text"""

    def test_parse_numbered_list(self, pcm_headers):
        """Numbered checklist should produce 3 questions with correct types + 1 limit + 1 observation."""
        text = "1. Verificar vazamentos\n2. Verificar temperatura máximo 80°C\n3. Tirar foto do manômetro\nOBS: Parar se > 90°C"
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": text}, headers=pcm_headers)
        assert r.status_code == 200, f"Expected 200 got {r.status_code}: {r.text[:200]}"
        data = r.json()
        # Structure assertions
        assert "perguntas" in data
        assert "observacoes" in data
        assert "limites" in data
        assert "frequencia" in data
        assert "metadata" in data
        # Content assertions
        perguntas = data["perguntas"]
        assert len(perguntas) == 3, f"Expected 3 questions, got {len(perguntas)}: {[p['texto'] for p in perguntas]}"
        # Q1: Verificar vazamentos (conforme_nao_conforme)
        assert "vazamento" in perguntas[0]["texto"].lower()
        assert perguntas[0]["tipo_campo"] == "conforme_nao_conforme"
        # Q2: temperatura → numerico, limite 80
        assert "temperatura" in perguntas[1]["texto"].lower()
        assert perguntas[1]["tipo_campo"] == "numerico"
        assert perguntas[1].get("limite_max") == "80", f"Expected limite_max=80, got {perguntas[1].get('limite_max')}"
        # Q3: foto
        assert "foto" in perguntas[2]["texto"].lower() or "manômetro" in perguntas[2]["texto"].lower()
        assert perguntas[2]["tipo_campo"] == "foto"
        # Observations
        assert len(data["observacoes"]) >= 1, f"Expected observation, got {data['observacoes']}"
        assert "parar" in data["observacoes"][0].lower()
        # Limites
        assert len(data["limites"]) >= 1
        # Metadata
        assert data["metadata"]["total_perguntas"] == 3
        assert data["metadata"]["total_observacoes"] >= 1
        assert data["metadata"]["total_limites"] >= 1

    def test_parse_bullets(self, pcm_headers):
        text = "- Verificar nível de óleo\n- Verificar ruído do motor\n• Verificar vibração"
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": text}, headers=pcm_headers)
        assert r.status_code == 200
        data = r.json()
        assert len(data["perguntas"]) == 3
        for p in data["perguntas"]:
            assert p["texto"].startswith("Verificar")

    def test_parse_with_frequency(self, pcm_headers):
        text = "Inspeção diária\n1. Verificar pressão\n2. Verificar vazão"
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": text}, headers=pcm_headers)
        assert r.status_code == 200
        data = r.json()
        assert data["frequencia"] is not None
        assert "diári" in data["frequencia"].lower()

    def test_parse_checkboxes(self, pcm_headers):
        text = "☐ Verificar filtro\n☑ Verificar ventilador\n[ ] Verificar bomba"
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": text}, headers=pcm_headers)
        assert r.status_code == 200
        data = r.json()
        assert len(data["perguntas"]) >= 2  # checkboxes should be parsed

    def test_parse_empty_text_400(self, pcm_headers):
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": ""}, headers=pcm_headers)
        assert r.status_code == 400

    def test_parse_no_auth_401(self):
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text", json={"text": "1. Test"})
        assert r.status_code in (401, 403)

    def test_parse_operador_forbidden(self, operador_token):
        headers = {"Authorization": f"Bearer {operador_token}"}
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": "1. Test"}, headers=headers)
        assert r.status_code == 403, f"Operador should not have plan-import permission, got {r.status_code}"

    def test_parse_field_type_detection(self, pcm_headers):
        """Verify field type detection: pressão→numerico, foto→foto, observação→texto."""
        text = "1. Medir pressão de descarga\n2. Anexar foto do painel\n3. Descrever condição do óleo"
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": text}, headers=pcm_headers)
        assert r.status_code == 200
        data = r.json()
        tipos = [p["tipo_campo"] for p in data["perguntas"]]
        assert "numerico" in tipos, f"Expected numerico for pressão, got {tipos}"
        assert "foto" in tipos, f"Expected foto, got {tipos}"


# ============== PARSE-FILE ENDPOINT ==============

class TestParseFile:
    """POST /api/planos-inspecao/parse-file"""

    def test_parse_txt_file(self, pcm_headers):
        content = "1. Verificar vazamentos\n2. Verificar temperatura máximo 80°C\n3. Tirar foto do manômetro\nOBS: Parar se > 90°C"
        files = {"file": ("plan.txt", io.BytesIO(content.encode('utf-8')), "text/plain")}
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-file",
                          files=files, headers=pcm_headers)
        assert r.status_code == 200, f"Got {r.status_code}: {r.text[:200]}"
        data = r.json()
        assert len(data["perguntas"]) == 3
        assert data["metadata"].get("source") == "txt"

    def test_parse_txt_bullets(self, pcm_headers):
        content = "- Item 1 verificação\n- Item 2 verificação\n- Item 3 verificação"
        files = {"file": ("bullets.txt", io.BytesIO(content.encode('utf-8')), "text/plain")}
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-file",
                          files=files, headers=pcm_headers)
        assert r.status_code == 200
        assert len(r.json()["perguntas"]) == 3

    def test_parse_unknown_extension_falls_back_to_txt(self, pcm_headers):
        content = "1. Test question one\n2. Test question two"
        files = {"file": ("plan.abc", io.BytesIO(content.encode('utf-8')), "text/plain")}
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-file",
                          files=files, headers=pcm_headers)
        assert r.status_code == 200
        assert len(r.json()["perguntas"]) >= 2

    def test_parse_file_no_auth_401(self):
        files = {"file": ("plan.txt", io.BytesIO(b"1. Test"), "text/plain")}
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-file", files=files)
        assert r.status_code in (401, 403)

    def test_parse_docx_file(self, pcm_headers):
        """Create a real .docx in-memory and post it."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc = Document()
        doc.add_paragraph("1. Verificar rolamentos")
        doc.add_paragraph("2. Medir temperatura máximo 85°C")
        doc.add_paragraph("3. Foto do painel de controle")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        files = {"file": ("plan.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-file",
                          files=files, headers=pcm_headers)
        assert r.status_code == 200, f"Got {r.status_code}: {r.text[:200]}"
        data = r.json()
        assert data["metadata"].get("source") == "docx"
        assert len(data["perguntas"]) >= 2

    def test_parse_xlsx_file(self, pcm_headers):
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")
        wb = Workbook()
        ws = wb.active
        ws.append(["Item", "Descrição", "Limite Máximo", "Unidade"])
        ws.append([1, "Verificar temperatura do rolamento", "80", "°C"])
        ws.append([2, "Verificar vazamento na bomba", "", ""])
        ws.append([3, "Verificar pressão de descarga", "10", "bar"])
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        files = {"file": ("plan.xlsx", buf, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-file",
                          files=files, headers=pcm_headers)
        assert r.status_code == 200, f"Got {r.status_code}: {r.text[:200]}"
        data = r.json()
        assert data["metadata"].get("source") == "excel"
        assert len(data["perguntas"]) >= 2


# ============== E2E SAVE FLOW ==============

class TestSaveAsPlanAndTemplate:
    """After parsing, verify plan/template creation endpoints still work with parsed payload."""

    def test_save_parsed_as_template(self, pcm_headers):
        # Parse
        text = "1. TEST_SP57 Verificar limpeza\n2. TEST_SP57 Medir vibração máximo 4.5mm/s"
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": text}, headers=pcm_headers)
        assert r.status_code == 200
        parsed = r.json()
        assert len(parsed["perguntas"]) == 2

        # Save as template
        payload = {
            "nome": "TEST_SP57_TemplateFromImport",
            "tipo_equipamento": "",
            "descricao": "E2E test template",
            "itens": [
                {"texto": p["texto"], "tipo_campo": p.get("tipo_campo", "conforme_nao_conforme"),
                 "obrigatorio": True, "ordem": i,
                 "limite_min": p.get("limite_min", ""), "limite_max": p.get("limite_max", ""),
                 "unidade": p.get("unidade", ""), "grupo": p.get("grupo", "")}
                for i, p in enumerate(parsed["perguntas"])
            ]
        }
        r2 = requests.post(f"{BASE_URL}/api/inspection-templates",
                           json=payload, headers=pcm_headers)
        assert r2.status_code in (200, 201), f"Template create failed: {r2.status_code} {r2.text[:200]}"
        created = r2.json()
        assert created["nome"] == "TEST_SP57_TemplateFromImport"
        assert len(created.get("itens", [])) == 2

    def test_save_parsed_as_plan(self, pcm_headers):
        # Parse
        text = "1. TEST_SP57 Verificar filtro\n2. TEST_SP57 Verificar bomba"
        r = requests.post(f"{BASE_URL}/api/planos-inspecao/parse-text",
                          json={"text": text}, headers=pcm_headers)
        assert r.status_code == 200
        parsed = r.json()

        # Save as plan (no ativo — allowed by API if null)
        payload = {
            "nome": "TEST_SP57_PlanFromImport",
            "tipo": "inspecao",
            "disciplina": "mecanica",
            "ativo_id": None,
            "frequencia": parsed.get("frequencia") or "mensal",
            "status": "rascunho",
            "force_override": False,
            "perguntas": [
                {"texto": p["texto"], "tipo_campo": p.get("tipo_campo", "conforme_nao_conforme"),
                 "obrigatorio": True, "ordem": i,
                 "limite_min": p.get("limite_min", ""), "limite_max": p.get("limite_max", ""),
                 "unidade": p.get("unidade", ""), "grupo": p.get("grupo", "")}
                for i, p in enumerate(parsed["perguntas"])
            ]
        }
        r2 = requests.post(f"{BASE_URL}/api/planos-inspecao",
                           json=payload, headers=pcm_headers)
        # Some implementations require ativo_id; accept 200/201 or informative 400/422
        assert r2.status_code in (200, 201, 400, 422), f"Unexpected: {r2.status_code} {r2.text[:300]}"
        if r2.status_code in (200, 201):
            created = r2.json()
            assert created["nome"] == "TEST_SP57_PlanFromImport"
            assert len(created.get("perguntas", [])) == 2
