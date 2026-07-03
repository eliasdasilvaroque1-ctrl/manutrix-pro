"""
MAINTRIX — Plan Import Parser
Smart parser for extracting checklist questions from text, PDF, Excel, Word, and TXT files.
No AI required — rule-based pattern recognition.
"""
import re
import io
import uuid
from typing import List, Dict, Optional


# ============== TEXT PARSER ==============

# Patterns that indicate a checklist item
LINE_PATTERNS = [
    # Numbered: 1. 2. 3) etc
    re.compile(r'^\s*(\d+)[.):\-]\s+(.+)'),
    # Lettered: a) b. c- etc
    re.compile(r'^\s*([a-zA-Z])[.):\-]\s+(.+)'),
    # Bullets: - • · ● ○ ▪ ■ ★
    re.compile(r'^\s*[\-•·●○▪■★►→]\s+(.+)'),
    # Checkboxes: ☐ □ ✓ ✔ ☑ [ ] [x] [X]
    re.compile(r'^\s*[☐□✓✔☑]\s*(.+)'),
    re.compile(r'^\s*\[[\sx✓XxV]?\]\s*(.+)'),
    # Tab/indent followed by text (common in pasted plans)
    re.compile(r'^\t+\s*(.+)'),
]

# Patterns that indicate a limit/threshold
LIMIT_PATTERN = re.compile(
    r'(m[áa]x(?:imo)?|m[íi]n(?:imo)?|limite|faixa|range|toler[âa]ncia|ideal|aceit[áa]vel)'
    r'[\s:=]*'
    r'([\d.,]+)\s*([°ºCcFf%BarbarPSIpsiMPampaKPakpaRPMrpmHzhz℃μmmmcmm²]*)',
    re.IGNORECASE
)

TEMPERATURE_PATTERN = re.compile(r'(\d+)\s*[°ºC℃]')
NUMERIC_LIMIT_PATTERN = re.compile(r'[<>≤≥]=?\s*(\d+[.,]?\d*)\s*(\w*)')

# Patterns for sections/headers (not questions)
HEADER_PATTERNS = [
    re.compile(r'^\s*(SEÇÃO|SECTION|ÁREA|AREA|GRUPO|GROUP|ETAPA|FASE|BLOCO)\s', re.IGNORECASE),
    re.compile(r'^[\s=\-_]{3,}$'),  # Separator lines
    re.compile(r'^\s*$'),  # Empty lines
]

OBSERVATION_KEYWORDS = ['obs', 'observ', 'nota', 'atenção', 'atencao', 'cuidado', 'aviso', 'warning', 'note']
FREQUENCY_KEYWORDS = ['diári', 'semanal', 'mensal', 'trimestral', 'semestral', 'anual', 'quinzenal', 'hora', 'turno']


def is_header_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return True
    if len(stripped) < 4:
        return True
    for p in HEADER_PATTERNS:
        if p.match(stripped):
            return True
    # All caps short line = likely header
    if stripped.isupper() and len(stripped) < 60:
        return True
    return False


def extract_limit(text: str) -> Optional[Dict]:
    """Try to extract numeric limits from text."""
    m = LIMIT_PATTERN.search(text)
    if m:
        return {"tipo": m.group(1).lower().strip(), "valor": m.group(2), "unidade": m.group(3) or ""}
    m = TEMPERATURE_PATTERN.search(text)
    if m:
        return {"tipo": "limite", "valor": m.group(1), "unidade": "°C"}
    m = NUMERIC_LIMIT_PATTERN.search(text)
    if m:
        return {"tipo": "limite", "valor": m.group(1), "unidade": m.group(2) or ""}
    return None


def detect_frequency(text: str) -> Optional[str]:
    t = text.lower()
    for kw in FREQUENCY_KEYWORDS:
        if kw in t:
            # Find the full word
            m = re.search(rf'\b\w*{kw}\w*\b', t, re.IGNORECASE)
            return m.group(0) if m else kw
    return None


def parse_text(text: str) -> Dict:
    """Parse plain text into structured checklist items."""
    lines = text.strip().split('\n')
    perguntas = []
    observacoes = []
    limites = []
    frequencia = None
    current_section = ""

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        # Check for frequency
        if not frequencia:
            freq = detect_frequency(line)
            if freq:
                frequencia = freq

        # Check if it's a header/section
        if is_header_line(line):
            # Check if it's a section header (useful context)
            if line.isupper() and len(line) > 3:
                current_section = line.title()
            continue

        # Check if it's an observation
        if any(kw in line.lower()[:20] for kw in OBSERVATION_KEYWORDS):
            observacoes.append(line)
            continue

        # Try to match as a checklist item
        matched = False
        question_text = line

        for pattern in LINE_PATTERNS:
            m = pattern.match(raw_line)
            if m:
                # Get the actual question text (last group)
                question_text = m.group(m.lastindex).strip()
                matched = True
                break

        # If no pattern matched but it's a reasonable line, still treat as question
        if not matched and len(line) > 10 and not line.startswith('#'):
            question_text = line
            matched = True

        if matched and question_text and len(question_text) > 3:
            # Detect type of response
            tipo_campo = detect_field_type(question_text)

            # Extract limits
            limit = extract_limit(question_text)
            pergunta = {
                "id": str(uuid.uuid4()),
                "texto": clean_question_text(question_text),
                "tipo_campo": tipo_campo,
                "obrigatorio": True,
                "ordem": len(perguntas),
                "grupo": current_section or "",
            }
            if limit:
                pergunta["limite_min"] = ""
                pergunta["limite_max"] = limit["valor"]
                pergunta["unidade"] = limit["unidade"]
                limites.append(limit)

            perguntas.append(pergunta)

    return {
        "perguntas": perguntas,
        "observacoes": observacoes,
        "limites": limites,
        "frequencia": frequencia,
        "metadata": {
            "total_perguntas": len(perguntas),
            "total_observacoes": len(observacoes),
            "total_limites": len(limites),
            "frequencia_detectada": frequencia,
        }
    }


def detect_field_type(text: str) -> str:
    t = text.lower()
    # Numeric indicators
    if any(kw in t for kw in ['temperatura', 'pressão', 'pressao', 'vazão', 'vazao', 'rpm', 'nivel', 'nível',
                                'valor', 'medição', 'medicao', 'corrente', 'tensão', 'tensao', 'vibração', 'vibracao',
                                'torque', 'carga', 'frequência', 'frequencia', 'potência', 'potencia']):
        return "numerico"
    if any(kw in t for kw in ['foto', 'imagem', 'evidência', 'evidencia', 'registro fotográfico', 'anexar']):
        return "foto"
    if any(kw in t for kw in ['observ', 'coment', 'descrever', 'detalhar', 'relatar', 'informar detalhes']):
        return "texto"
    # Default: conforme/não conforme
    return "conforme_nao_conforme"


def clean_question_text(text: str) -> str:
    """Clean up question text."""
    # Remove trailing colons, periods if it's just a label
    text = text.strip()
    # Remove check markers at start
    text = re.sub(r'^[\[\]☐□✓✔☑xX\s]+', '', text).strip()
    # Capitalize first letter
    if text and text[0].islower():
        text = text[0].upper() + text[1:]
    return text


# ============== FILE PARSERS ==============

def parse_pdf(file_bytes: bytes) -> Dict:
    """Extract text from PDF and parse it."""
    import pdfplumber
    text_parts = []
    tables = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # Extract tables first
            page_tables = page.extract_tables()
            for table in page_tables:
                tables.extend(table)
            # Extract text
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = '\n'.join(text_parts)

    # Parse tables into questions too
    table_perguntas = []
    for row in tables:
        if not row or not any(row):
            continue
        cells = [str(c).strip() for c in row if c]
        if not cells:
            continue
        # Skip header-like rows
        if any(h in cells[0].lower() for h in ['item', 'nº', 'n°', '#', 'seq']):
            continue
        # First substantive cell is the question
        main_text = cells[0] if len(cells) > 0 else ''
        if len(main_text) < 4:
            main_text = ' '.join(cells[:2]) if len(cells) > 1 else main_text
        if main_text and len(main_text) > 3:
            p = {
                "id": str(uuid.uuid4()),
                "texto": clean_question_text(main_text),
                "tipo_campo": detect_field_type(main_text),
                "obrigatorio": True,
                "ordem": len(table_perguntas),
                "grupo": "",
            }
            # Check other cells for limits
            for cell in cells[1:]:
                limit = extract_limit(cell)
                if limit:
                    p["limite_max"] = limit["valor"]
                    p["unidade"] = limit.get("unidade", "")
            table_perguntas.append(p)

    # Parse text
    result = parse_text(full_text)

    # Merge table questions (avoid duplicates)
    existing_texts = {p['texto'].lower() for p in result['perguntas']}
    for tp in table_perguntas:
        if tp['texto'].lower() not in existing_texts:
            tp['ordem'] = len(result['perguntas'])
            result['perguntas'].append(tp)
            existing_texts.add(tp['texto'].lower())

    result['metadata']['total_perguntas'] = len(result['perguntas'])
    result['metadata']['source'] = 'pdf'
    return result


def parse_excel(file_bytes: bytes) -> Dict:
    """Parse Excel file into questions."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    perguntas = []
    limites = []

    for ws in wb.worksheets:
        header_row = None
        desc_col = None
        limit_col = None
        unit_col = None

        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(c).strip() if c else '' for c in row]
            if not any(cells):
                continue

            # Detect header row
            if header_row is None:
                lower_cells = [c.lower() for c in cells]
                if any(h in ' '.join(lower_cells) for h in ['descrição', 'descricao', 'item', 'pergunta', 'verificar', 'check']):
                    header_row = row_idx
                    for i, c in enumerate(lower_cells):
                        if any(h in c for h in ['descrição', 'descricao', 'item', 'pergunta', 'verificar', 'check', 'atividade']):
                            desc_col = i
                        if any(h in c for h in ['limite', 'max', 'min', 'faixa', 'range', 'valor']):
                            limit_col = i
                        if any(h in c for h in ['unidade', 'unit', 'und']):
                            unit_col = i
                    continue

            # Data rows
            main_text = ''
            if desc_col is not None and desc_col < len(cells):
                main_text = cells[desc_col]
            else:
                # Use first non-empty, non-numeric cell
                for c in cells:
                    if c and len(c) > 3 and not c.replace('.', '').replace(',', '').isdigit():
                        main_text = c
                        break

            if main_text and len(main_text) > 3:
                p = {
                    "id": str(uuid.uuid4()),
                    "texto": clean_question_text(main_text),
                    "tipo_campo": detect_field_type(main_text),
                    "obrigatorio": True,
                    "ordem": len(perguntas),
                    "grupo": "",
                }
                if limit_col is not None and limit_col < len(cells) and cells[limit_col]:
                    limit = extract_limit(cells[limit_col])
                    if limit:
                        p["limite_max"] = limit["valor"]
                        limites.append(limit)
                    else:
                        p["limite_max"] = cells[limit_col]
                if unit_col is not None and unit_col < len(cells) and cells[unit_col]:
                    p["unidade"] = cells[unit_col]

                perguntas.append(p)

    wb.close()
    return {
        "perguntas": perguntas,
        "observacoes": [],
        "limites": limites,
        "frequencia": None,
        "metadata": {"total_perguntas": len(perguntas), "total_limites": len(limites), "source": "excel"}
    }


def parse_docx(file_bytes: bytes) -> Dict:
    """Parse Word document into questions."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it's a list item
            if para.style and 'list' in (para.style.name or '').lower():
                text_parts.append(f"• {text}")
            else:
                text_parts.append(text)

    # Also parse tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append('\t'.join(cells))

    full_text = '\n'.join(text_parts)
    result = parse_text(full_text)
    result['metadata']['source'] = 'docx'
    return result


def parse_txt(file_bytes: bytes) -> Dict:
    """Parse plain text file."""
    # Try different encodings
    for enc in ['utf-8', 'latin-1', 'cp1252']:
        try:
            text = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = file_bytes.decode('utf-8', errors='replace')

    result = parse_text(text)
    result['metadata']['source'] = 'txt'
    return result
